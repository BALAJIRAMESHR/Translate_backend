from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel, Field
from typing import Optional
import google.generativeai as genai
import speech_recognition as sr
from googletrans import Translator
from typing import Optional, Tuple
import pyttsx3
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import tempfile
import logging
from enum import Enum
import io
import os
import wave
import soundfile as sf
from pydub import AudioSegment
import numpy as np
import speech_recognition as sr
from tempfile import SpooledTemporaryFile
import shutil
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import base64
from PIL import Image
import uvicorn

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

# Supported languages list
SUPPORTED_LANGUAGES = {
    "en": "English",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "it": "Italian",
    "pt": "Portuguese",
    "ru": "Russian",
    "ja": "Japanese",
    "ko": "Korean",
    "zh": "Chinese",
    "ar": "Arabic",
    "hi": "Hindi",
    "ta": "Tamil",
    "kn": "Kannada",
    "te": "Telugu",
}

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_IMAGE_SIZE = 5 * 1024 * 1024  # 5MB limit for images
SUPPORTED_IMAGE_FORMATS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}
# Initialize translation components
translator = Translator()
recognizer = sr.Recognizer()


class TranslationRequest(BaseModel):
    text: str = Field(..., min_length=1, max_length=5000)
    target_language: str
    source_language: Optional[str] = None


class TranslationResponse(BaseModel):
    translated_text: str
    original_text: str
    target_language: str
    source_language: Optional[str]
    detected_language: Optional[str] = None


class AudioTranslationResponse(BaseModel):
    detected_text: str
    translated_text: str
    source_language: str
    target_language: str


class SupportedFileType(str, Enum):
    TXT = "txt"
    DOCX = "docx"
    PDF = "pdf"


VALID_CONTENT_TYPES = {
    SupportedFileType.TXT: "text/plain",
    SupportedFileType.DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    SupportedFileType.PDF: "application/pdf",
}

app = FastAPI(
    title="Multilingual Translation Service",
    description="API for text, file, and voice translation between multiple languages",
    version="1.0.0",
)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure Gemini API
GEMINI_API_KEY = "AIzaSyAnEbQDilqTRSE1Bn8dqAVNhf6Ml_YyX18"
genai.configure(api_key=GEMINI_API_KEY)


async def text_to_speech(text: str, language: str) -> bytes:
    """
    Convert text to speech using pyttsx3
    """
    try:
        # Create a temporary file with context manager
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_file:
            temp_path = temp_file.name

            engine = pyttsx3.init()
            engine.setProperty("rate", 150)
            engine.setProperty("volume", 0.9)

            # Set language if available
            voices = engine.getProperty("voices")
            for voice in voices:
                if language in voice.languages:
                    engine.setProperty("voice", voice.id)
                    break

            # Save speech to temporary file
            engine.save_to_file(text, temp_path)
            engine.runAndWait()

            # Read the generated audio file
            with open(temp_path, "rb") as audio_file:
                audio_content = audio_file.read()

            return audio_content

    except Exception as e:
        logger.error(f"Text-to-speech error: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Text-to-speech conversion failed: {str(e)}"
        )
    finally:
        # Clean up temporary file
        if os.path.exists(temp_path):
            os.unlink(temp_path)


async def translate_text_with_gemini(
    text: str, target_language: str, source_language: Optional[str] = None
) -> str:
    try:
        model = genai.GenerativeModel("gemini-pro")

        # Construct the translation prompt
        if source_language:
            prompt = f"""
            Translate the following text from {SUPPORTED_LANGUAGES[source_language]} to {SUPPORTED_LANGUAGES[target_language]}.
            Preserve all formatting, special characters, and maintain the original meaning.
            
            Text to translate:
            {text}
            """
        else:
            prompt = f"""
            Translate the following text to {SUPPORTED_LANGUAGES[target_language]}.
            Detect the source language, preserve all formatting, special characters, and maintain the original meaning.
            
            Text to translate:
            {text}
            """

        response = model.generate_content(prompt)
        if not response or not response.text:
            raise HTTPException(
                status_code=500, detail="Empty response from translation service"
            )
        return response.text.strip()
    except Exception as e:
        logger.error(f"Translation error: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Translation service failed: {str(e)}"
        )


async def validate_file(
    file: UploadFile, target_language: str, source_language: Optional[str] = None
) -> None:
    """
    Validate uploaded file format, size, and language parameters.
    """
    # Validate file presence
    if not file:
        raise HTTPException(status_code=400, detail="No file provided")

    # Validate file size (10MB limit)
    content = await file.read()
    await file.seek(0)  # Reset file pointer
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File size exceeds maximum limit of {MAX_FILE_SIZE/1024/1024}MB",
        )

    # Validate file extension
    file_ext = file.filename.split(".")[-1].lower()
    if file_ext not in [e.value for e in SupportedFileType]:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format. Supported formats: {', '.join([e.value for e in SupportedFileType])}",
        )

    # Validate content type
    if file.content_type not in VALID_CONTENT_TYPES.values():
        raise HTTPException(
            status_code=400,
            detail=f"Invalid content type. Supported types: {', '.join(VALID_CONTENT_TYPES.values())}",
        )

    # Validate target language
    if target_language not in SUPPORTED_LANGUAGES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported target language. Supported languages: {', '.join(SUPPORTED_LANGUAGES.keys())}",
        )

    # Validate source language if provided
    if source_language and source_language not in SUPPORTED_LANGUAGES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported source language. Supported languages: {', '.join(SUPPORTED_LANGUAGES.keys())}",
        )


async def process_file_content(
    file_path: str,
    file_type: SupportedFileType,
    target_language: str,
    source_language: Optional[str] = None,
) -> tuple[str, bytes]:
    try:
        if file_type == SupportedFileType.TXT:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            translated_content = await translate_text_with_gemini(
                content, target_language, source_language
            )
            return "txt", translated_content.encode("utf-8")

        elif file_type == SupportedFileType.DOCX:
            doc = Document(file_path)
            new_doc = Document()

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    translated_text = await translate_text_with_gemini(
                        paragraph.text, target_language, source_language
                    )
                    new_paragraph = new_doc.add_paragraph()
                    new_paragraph.style = paragraph.style
                    new_paragraph.paragraph_format.alignment = (
                        paragraph.paragraph_format.alignment
                    )
                    new_paragraph.text = translated_text

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                new_doc.save(tmp.name)
                with open(tmp.name, "rb") as f:
                    content = f.read()
            os.unlink(tmp.name)
            return "docx", content

        elif file_type == SupportedFileType.PDF:
            reader = PdfReader(file_path)
            translated_pages = []

            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    translated_text = await translate_text_with_gemini(
                        text, target_language, source_language
                    )
                    translated_pages.append(translated_text)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                c = canvas.Canvas(tmp.name, pagesize=letter)
                width, height = letter

                for page_text in translated_pages:
                    y = height - 50
                    for line in page_text.split("\n"):
                        if y < 50:
                            c.showPage()
                            y = height - 50
                        c.drawString(50, y, line[:100])
                        y -= 20
                    c.showPage()
                c.save()

                with open(tmp.name, "rb") as f:
                    content = f.read()
            os.unlink(tmp.name)
            return "pdf", content

    except Exception as e:
        logger.error(f"File processing error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"File processing failed: {str(e)}")


async def validate_wav_file(content: bytes) -> bool:
    """Validate if the content is a proper WAV file"""
    try:
        with io.BytesIO(content) as wav_io:
            with wave.open(wav_io, "rb") as wave_read:
                # Check basic WAV properties
                channels = wave_read.getnchannels()
                sample_width = wave_read.getsampwidth()
                framerate = wave_read.getframerate()

                # Basic validation
                if channels not in [1, 2]:  # Mono or Stereo
                    return False
                if sample_width not in [1, 2, 4]:  # 8, 16, or 32 bit
                    return False
                if framerate < 8000 or framerate > 48000:  # Common sampling rates
                    return False

        return True
    except Exception:
        return False


async def process_voice(
    audio_file: UploadFile, target_language: str
) -> tuple[str, str]:
    """Process voice file and translate the detected text with improved Windows compatibility."""
    recognizer = sr.Recognizer()

    try:
        # Create SpooledTemporaryFile objects to handle the audio data in memory first
        with SpooledTemporaryFile() as original_audio, SpooledTemporaryFile() as converted_audio:
            # Read the uploaded content
            content = await audio_file.read()
            original_audio.write(content)
            original_audio.seek(0)

            try:
                # Convert to WAV using pydub without saving to disk first
                audio = AudioSegment.from_file(original_audio)

                # Process audio
                audio = audio.set_channels(1)  # Convert to mono
                audio = audio.set_frame_rate(16000)  # Set sample rate to 16kHz

                # Export to the second temporary file
                audio.export(
                    converted_audio,
                    format="wav",
                    parameters=[
                        "-acodec",
                        "pcm_s16le",  # Use PCM 16-bit encoding
                    ],
                )
                converted_audio.seek(0)

                # Create a temporary WAV file just for speech recognition
                with tempfile.NamedTemporaryFile(
                    suffix=".wav", delete=False
                ) as temp_wav:
                    temp_wav_path = temp_wav.name
                    shutil.copyfileobj(converted_audio, temp_wav)

                try:
                    # Use speech recognition on the temporary WAV file
                    with sr.AudioFile(temp_wav_path) as source:
                        # Adjust for ambient noise
                        recognizer.adjust_for_ambient_noise(source, duration=0.5)
                        audio_data = recognizer.record(source)

                        # Attempt speech recognition
                        try:
                            detected_text = recognizer.recognize_google(
                                audio_data,
                                language="en-US",  # Default to English
                                show_all=False,
                            )

                            if not detected_text:
                                raise HTTPException(
                                    status_code=400,
                                    detail="No speech detected in the audio file",
                                )

                            # Translate detected text
                            translated_text = await translate_text_with_gemini(
                                detected_text, target_language
                            )

                            return detected_text, translated_text

                        except sr.UnknownValueError:
                            raise HTTPException(
                                status_code=400,
                                detail="Speech could not be recognized. Please ensure clear audio with minimal background noise.",
                            )
                        except sr.RequestError as e:
                            raise HTTPException(
                                status_code=500,
                                detail=f"Speech recognition service error: {str(e)}",
                            )

                finally:
                    # Clean up the temporary WAV file
                    if os.path.exists(temp_wav_path):
                        os.unlink(temp_wav_path)

            except Exception as e:
                logger.error(f"Audio processing error: {str(e)}")
                raise HTTPException(
                    status_code=500, detail=f"Audio conversion error: {str(e)}"
                )

    except Exception as e:
        logger.error(f"Voice processing error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


async def validate_image(file: UploadFile) -> None:
    """Validate image file format and size"""
    # Check file presence
    if not file:
        raise HTTPException(status_code=400, detail="No image file provided")

    # Validate file extension
    if not any(file.filename.lower().endswith(fmt) for fmt in SUPPORTED_IMAGE_FORMATS):
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported image format. Supported formats: {', '.join(SUPPORTED_IMAGE_FORMATS)}",
        )

    # Read content and validate size
    content = await file.read()
    await file.seek(0)

    if len(content) > MAX_IMAGE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"Image size exceeds maximum limit of {MAX_IMAGE_SIZE/1024/1024}MB",
        )

    # Validate image can be opened/processed
    try:
        Image.open(io.BytesIO(content))
    except Exception:
        raise HTTPException(
            status_code=400, detail="Invalid image file. Could not process the image."
        )


async def extract_and_translate_text_with_gemini(
    image_content: bytes, target_language: str, source_language: Optional[str] = None
) -> tuple[str, str]:
    """Extract and translate text from image using Gemini Vision"""
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")

        # Create a PIL Image from bytes
        image = Image.open(io.BytesIO(image_content))

        # Extract text prompt
        extract_prompt = "Extract all the text from this image. Return only the extracted text, nothing else."

        # Generate content using the image directly
        extract_response = model.generate_content([extract_prompt, image])

        if not extract_response or not extract_response.text:
            raise HTTPException(status_code=400, detail="No text detected in the image")

        detected_text = extract_response.text.strip()

        # Use gemini-pro model for translation
        translation_model = genai.GenerativeModel("gemini-pro")

        # Translate the detected text
        if source_language:
            translate_prompt = f"""
            Translate the following text from {SUPPORTED_LANGUAGES[source_language]} to {SUPPORTED_LANGUAGES[target_language]}.
            Preserve all formatting and maintain the original meaning.
            
            Text to translate:
            {detected_text}
            """
        else:
            translate_prompt = f"""
            Translate the following text to {SUPPORTED_LANGUAGES[target_language]}.
            Preserve all formatting and maintain the original meaning.
            
            Text to translate:
            {detected_text}
            """

        translate_response = translation_model.generate_content(translate_prompt)
        if not translate_response or not translate_response.text:
            raise HTTPException(status_code=500, detail="Translation failed")

        translated_text = translate_response.text.strip()

        return detected_text, translated_text

    except Exception as e:
        logger.error(f"Gemini processing error: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Failed to process image: {str(e)}"
        )


class ImageTranslationResponse(BaseModel):
    detected_text: str
    translated_text: str
    source_language: Optional[str]
    target_language: str


# API Endpoints
@app.get("/api/languages")
async def get_supported_languages():
    """Get list of supported languages"""
    return SUPPORTED_LANGUAGES


@app.post("/api/translate", response_model=TranslationResponse)
async def translate_text(request: TranslationRequest):
    try:
        if not request.text.strip():
            raise HTTPException(
                status_code=400, detail="Text to translate cannot be empty"
            )

        if request.target_language not in SUPPORTED_LANGUAGES:
            raise HTTPException(status_code=400, detail="Unsupported target language")

        if (
            request.source_language
            and request.source_language not in SUPPORTED_LANGUAGES
        ):
            raise HTTPException(status_code=400, detail="Unsupported source language")

        translated_text = await translate_text_with_gemini(
            request.text, request.target_language, request.source_language
        )

        return TranslationResponse(
            translated_text=translated_text,
            original_text=request.text,
            target_language=request.target_language,
            source_language=request.source_language,
        )

    except Exception as e:
        logger.error(f"Translation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")


@app.post("/api/translate/file")
async def translate_file(
    file: UploadFile = File(...),
    target_language: str = Form(...),
    source_language: Optional[str] = Form(None),
):
    try:
        # Validate the file and parameters
        await validate_file(file, target_language, source_language)

        content = await file.read()
        file_ext = file.filename.split(".")[-1].lower()

        # Create temp directory if it doesn't exist
        temp_dir = "temp"
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        temp_input_path = os.path.join(temp_dir, f"input_{file.filename}")

        try:
            # Write content to temporary file
            with open(temp_input_path, "wb") as f:
                f.write(content)

            # Process the file
            file_type = SupportedFileType(file_ext)
            content_type, translated_content = await process_file_content(
                temp_input_path, file_type, target_language, source_language
            )

            # Prepare the response
            output = io.BytesIO(translated_content)
            media_type = VALID_CONTENT_TYPES[file_type]

            return StreamingResponse(
                output,
                media_type=media_type,
                headers={
                    "Content-Disposition": f"attachment; filename=translated_{file.filename}"
                },
            )

        finally:
            # Clean up temporary file
            if os.path.exists(temp_input_path):
                os.remove(temp_input_path)

    except HTTPException as http_ex:
        raise http_ex
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/translate/voice")
async def translate_voice(
    audio_file: UploadFile = File(...), target_language: str = Form(...)
):
    """
    Endpoint to handle voice translation requests with improved error handling.
    """
    # Validate file size
    content = await audio_file.read()
    if len(content) > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(status_code=400, detail="File size exceeds 10MB limit")

    # Reset file pointer for processing
    await audio_file.seek(0)

    try:
        detected_text, translated_text = await process_voice(
            audio_file, target_language
        )

        return JSONResponse(
            {
                "detected_text": detected_text,
                "translated_text": translated_text,
                "source_language": "auto-detected",
                "target_language": target_language,
            }
        )

    except HTTPException as http_ex:
        raise http_ex
    except Exception as e:
        logger.error(f"Voice translation error: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Voice translation failed: {str(e)}"
        )


@app.post("/api/translate/image", response_model=ImageTranslationResponse)
async def translate_image(
    image: UploadFile = File(...),
    target_language: str = Form(...),
    source_language: Optional[str] = Form(None),
):
    """
    Endpoint to handle image translation requests.
    Extracts text from images and translates it to the target language.
    """
    try:
        # Validate image and parameters
        await validate_image(image)

        if target_language not in SUPPORTED_LANGUAGES:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported target language. Supported languages: {', '.join(SUPPORTED_LANGUAGES.keys())}",
            )

        if source_language and source_language not in SUPPORTED_LANGUAGES:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported source language. Supported languages: {', '.join(SUPPORTED_LANGUAGES.keys())}",
            )

        # Read image content
        content = await image.read()

        # Extract and translate text
        detected_text, translated_text = await extract_and_translate_text_with_gemini(
            content, target_language, source_language
        )

        return ImageTranslationResponse(
            detected_text=detected_text,
            translated_text=translated_text,
            source_language=source_language or "auto-detected",
            target_language=target_language,
        )

    except HTTPException as http_ex:
        raise http_ex
    except Exception as e:
        logger.error(f"Image translation error: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Image translation failed: {str(e)}"
        )


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", 8000)))
